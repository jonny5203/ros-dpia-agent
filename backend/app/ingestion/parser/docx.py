from __future__ import annotations

import io
import re

from app.ingestion.parser.base import Parser
from app.ingestion.types import ParsedDocument, ParsedSection


class DocxParser(Parser):
    name = "python-docx"
    version = "python-docx>=1.1"
    _HEADING = re.compile(r"^Heading\s*(\d+)$")

    def parse(self, data: bytes, filename: str) -> ParsedDocument:
        from docx import Document
        doc = Document(io.BytesIO(data))
        sections: list[ParsedSection] = []
        current_path: list[str] = []

        for para in doc.paragraphs:
            style = para.style.name if para.style is not None else ""
            m = self._HEADING.match(style)
            text = para.text.strip()

            if not text:
                continue

            if m:
                level = int(m.group(1))
                current_path = [*current_path[: level - 1], text]
                continue
            sections.append(ParsedSection(
                text=text,
                page=None,
                section_title=current_path[-1] if current_path else None,
                section_path=" > ".join(current_path) or None,
            ))

        return ParsedDocument(
            sections=sections,
            parser=self.name,
            parser_version=self.version
        )
